"""
unified_agent/rag/file_parser.py
知识库文件解析器: 从上传的 pdf/docx/txt/md 文件提取纯文本.

设计说明 (D2 文件上传管控, 决策 2: Python 端解析):
- 前端上传 MultipartFile → Java 落盘 + 转发原始字节 → Python 本模块解析 → 返回 text;
- Java 作为 content SSOT, 拿到 text 后走原有 publish→kb_sync 链路 (content 写入 file_path);
- 解析在 Python 端执行的原因: pdf/docx 解析库 (pdfplumber/python-docx) 为 Python 生态,
  避免在 Java 引入重型解析依赖; 同时复用 Python 侧已有的 RAG 文件管理目录.

支持的文件类型与解析引擎:
- .pdf  → pdfplumber (支持文本+表格抽取, 扫描件 PDF 无文本层时返回空并标注);
- .docx → python-docx (按段落抽取, 含表格文本拼接);
- .txt/.md → 直读 UTF-8 文本 (Markdown 不转义, 保留原始语法供结构化分块).

容错策略:
- 文件损坏 / 格式不符 → 抛 ParseError (含面向用户的友好消息, 不泄漏堆栈);
- 解析库未安装 → 抛 ParseError 提示需安装依赖 (降级而非崩溃);
- 空文件 / 空文本 → 返回空字符串 + page_count=0, 由调用方决定是否拒绝.
"""
from __future__ import annotations

from dataclasses import dataclass
from typing import Optional

from core.logger import get_logger

logger = get_logger("kb_file_parser")

# 允许解析的扩展名白名单 (与 Java 侧 kb.upload.allowed-types 对齐)
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class ParseError(Exception):
    """文件解析异常 (面向用户的友好消息, 技术细节仅入日志).

    消失的库 / 损坏文件 / 不支持的格式统一抛此异常, 由路由层捕获后返回 200 + ok=false.
    """


@dataclass
class ParseResult:
    """文件解析结果.

    用 dataclass 而非 dict, 字段显式化便于调用方取值与类型检查.
    """

    # 提取的纯文本 (段落间以空行分隔, 供后续分块器处理)
    text: str
    # 页数 (PDF 为页数, docx/txt 为 1)
    page_count: int
    # 解析引擎名 (pdfplumber/python-docx/plain_text), 供日志与监控
    parse_engine: str
    # 原始文件名 (透传, 供日志关联)
    filename: str
    # 字符数 (text 长度, 供调用方做大小校验)
    char_count: int


def parse_file(filename: str, content: bytes) -> ParseResult:
    """按扩展名分发解析, 返回 ParseResult.

    Args:
        filename: 原始文件名 (用于推断扩展名与日志关联).
        content: 文件原始字节 (由 Java 转发或前端直传).

    Returns:
        ParseResult: 含 text/page_count/parse_engine.

    Raises:
        ParseError: 格式不支持 / 文件损坏 / 依赖缺失.
    """
    ext = _get_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        # 不支持的格式: 列出允许类型, 帮助用户纠正
        raise ParseError(f"不支持的文件类型: {ext}, 仅支持 {', '.join(sorted(ALLOWED_EXTENSIONS))}")

    try:
        if ext == ".pdf":
            return _parse_pdf(filename, content)
        if ext == ".docx":
            return _parse_docx(filename, content)
        # .txt / .md
        return _parse_plain(filename, content, ext)
    except ParseError:
        raise
    except Exception as exc:  # noqa: BLE001
        # 解析库抛出的底层异常统一包装为面向用户的友好消息 (技术细节仅入日志)
        logger.error(f"file_parse_failed filename={filename} ext={ext} error={exc}", exc_info=True)
        raise ParseError(f"文件解析失败, 请检查文件是否损坏或格式不符: {filename}")


def _get_extension(filename: str) -> str:
    """取小写扩展名 (含点), 无扩展名返回空串."""
    if not filename or "." not in filename:
        return ""
    return "." + filename.rsplit(".", 1)[-1].lower()


def _parse_pdf(filename: str, content: bytes) -> ParseResult:
    """用 pdfplumber 解析 PDF, 逐页抽取文本 (含表格文本).

    扫描件 PDF (无文本层) 抽取结果为空, 此时 page_count>0 但 text 为空,
    由调用方决定是否拒绝 (知识库不接受无文本的扫描件).
    """
    try:
        import pdfplumber
    except ImportError:
        raise ParseError("PDF 解析依赖未安装, 请执行 pip install pdfplumber")

    import io
    page_texts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            # extract_text 返回页面文本 (含表格被还原的文本); None 时降级为空串
            page_text = page.extract_text() or ""
            page_texts.append(page_text)
    text = "\n\n".join(page_texts).strip()
    logger.info(f"pdf_parsed filename={filename} pages={len(page_texts)} chars={len(text)}")
    return ParseResult(
        text=text,
        page_count=len(page_texts),
        parse_engine="pdfplumber",
        filename=filename,
        char_count=len(text),
    )


def _parse_docx(filename: str, content: bytes) -> ParseResult:
    """用 python-docx 解析 Word 文档, 按段落抽取 + 表格文本拼接.

    docx 表格单元格文本按行拼接为 | 分隔的文本行, 保留表格语义供后续表格感知分块.
    """
    try:
        import docx
    except ImportError:
        raise ParseError("Word 解析依赖未安装, 请执行 pip install python-docx")

    import io
    doc = docx.Document(io.BytesIO(content))
    parts = []
    # 段落文本 (正文)
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            parts.append(text)
    # 表格文本 (逐行逐格拼接, 保留 Markdown 表格语法供表格感知分块识别)
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                # 拼成 | a | b | c | 格式, 与 Markdown 表格行一致 (splitter._TABLE_LINE 可识别)
                parts.append("| " + " | ".join(cells) + " |")
    text = "\n\n".join(parts).strip()
    logger.info(f"docx_parsed filename={filename} chars={len(text)}")
    return ParseResult(
        text=text,
        page_count=1,
        parse_engine="python-docx",
        filename=filename,
        char_count=len(text),
    )


def _parse_plain(filename: str, content: bytes, ext: str) -> ParseResult:
    """直读 txt/md 文件 (UTF-8, 兼容 GBK 降级).

    Markdown 不做转义, 保留原始语法 (标题#/表格|/列表-) 供结构化分块器利用.
    """
    text = _decode_text(content)
    logger.info(f"plain_parsed filename={filename} ext={ext} chars={len(text)}")
    return ParseResult(
        text=text,
        page_count=1,
        parse_engine="plain_text",
        filename=filename,
        char_count=len(text),
    )


def _decode_text(content: bytes) -> str:
    """解码字节为文本: 优先 UTF-8, 失败回退 GBK (兼容 Windows 记事本默认编码)."""
    for encoding in ("utf-8", "gbk"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    # 两种编码均失败: 用 utf-8 忽略错误字节降级 (不阻断流程, 损失少量乱码字符)
    return content.decode("utf-8", errors="ignore").strip()
